"""Dual-format Word renderer for briefing packs.

Takes one or more verified briefing packs and produces a single .docx with
two presentations of the same verified material:

- Part A, whole paragraphs: the points composed into prose by a final LLM
  pass. The quotes are never given to the LLM as editable text; it writes
  connecting sentences around placeholder tokens ({{Q1}}, ...) and the
  verbatim quote plus its citation is substituted in mechanically. If the
  composition loses or duplicates a placeholder, the renderer falls back to a
  deterministic composition, so no verified point can be silently dropped and
  no quote can be altered.

- Part B, bullet points: one bullet per verified point, again with the
  verbatim quote and the WHO-format citation (IrisItem.citation()).

Run with:
    uv run --with python-docx python -m who_speech.render
(the demo builds a two-point pack from hand-verified points and renders it
to /tmp/who_render_demo.docx).
"""
from __future__ import annotations

import datetime
import re
from typing import TYPE_CHECKING

from pydantic import BaseModel, ConfigDict, Field

from who_speech import prompts

if TYPE_CHECKING:
    from who_speech.swarm import BriefingPack, BriefingPoint


class ComposedProse(BaseModel):
    model_config = ConfigDict(extra="forbid")
    paragraphs: list[str] = Field(default_factory=list)


def _placeholder(i: int) -> str:
    return f"{{{{Q{i + 1}}}}}"


def compose_paragraphs(pack: "BriefingPack") -> list[str]:
    """LLM pass: connecting prose around placeholder tokens, one per point.

    Returns paragraphs still containing the placeholders. Falls back to a
    deterministic one-point-per-paragraph composition when the LLM output
    does not use every placeholder exactly once.
    """
    from agents.tools.llm import StructuredOutputError, call_for_structured

    listed = "\n".join(
        f"[{_placeholder(i)}] {p.point}" for i, p in enumerate(pack.points)
    )
    user = f"Question: {pack.query}\n\nVerified points:\n{listed}"
    try:
        prose, _ = call_for_structured(
            system=prompts.COMPOSER_SYSTEM,
            user_message=user,
            output_schema=ComposedProse,
            usage_context="who_speech:composer",
            max_tokens=1200,
        )
        paragraphs = [p.strip() for p in prose.paragraphs if p.strip()]
    except StructuredOutputError:
        paragraphs = []

    joined = "\n".join(paragraphs)
    ok = paragraphs and all(
        joined.count(_placeholder(i)) == 1 for i in range(len(pack.points))
    )
    if ok:
        return paragraphs
    # Deterministic fallback: every point gets its own paragraph.
    return [f"{p.point} {_placeholder(i)}" for i, p in enumerate(pack.points)]


def _add_quote_runs(paragraph, point: "BriefingPoint") -> None:
    """Append the verbatim quote (italic, in quotation marks) + citation."""
    quote_run = paragraph.add_run(f"“{point.quote}”")
    quote_run.italic = True
    cite = f" ({point.citation}, p.{point.page})"
    paragraph.add_run(cite)


def _render_part_a(doc, pack: "BriefingPack", paragraphs: list[str]) -> None:
    token_to_point = {_placeholder(i): p for i, p in enumerate(pack.points)}
    splitter = re.compile(r"(\{\{Q\d+\}\})")
    for para_text in paragraphs:
        para = doc.add_paragraph()
        for piece in splitter.split(para_text):
            if not piece:
                continue
            point = token_to_point.get(piece)
            if point is not None:
                _add_quote_runs(para, point)
            else:
                para.add_run(piece)


def _render_part_b(doc, pack: "BriefingPack") -> None:
    for point in pack.points:
        bullet = doc.add_paragraph(style="List Bullet")
        lead = bullet.add_run(point.point)
        lead.bold = True
        bullet.add_run(" ")
        _add_quote_runs(bullet, point)


def render_briefing(
    packs: "list[BriefingPack] | BriefingPack",
    out_path: str,
    *,
    title: str = "WHO/Europe briefing: verified speaking points",
) -> str:
    """Render pack(s) to a single dual-format .docx and return the path."""
    from docx import Document
    from docx.shared import Pt

    if not isinstance(packs, list):
        packs = [packs]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, level=0)
    today = datetime.date.today().strftime("%d %B %Y")
    doc.add_paragraph(
        f"Prepared {today} from WHO/Europe publications in the IRIS "
        "repository. Every quotation is verbatim from the cited document, "
        "which carries a Creative Commons licence; each was checked "
        "mechanically against the source text before inclusion. Connecting "
        "prose in Part A restates the verified points only."
    )

    for pack in packs:
        doc.add_heading(pack.query, level=1)
        if pack.abstained or not pack.points:
            doc.add_paragraph(
                "No sufficiently supported points were found for this "
                "question; the system abstains rather than speculate."
                + (f" ({pack.note})" if pack.note else "")
            )
            continue

        doc.add_heading("Part A. Narrative", level=2)
        _render_part_a(doc, pack, compose_paragraphs(pack))

        doc.add_heading("Part B. Bullet points", level=2)
        _render_part_b(doc, pack)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    # Demo on two hand-verified Tajikistan points (see VERIFICATION_LOG.md).
    from who_speech.swarm import BriefingPack, BriefingPoint

    _cite = (
        "A situation assessment of rehabilitation in Tajikistan. "
        "World Health Organization; 2025. Licence: CC BY-NC-SA 3.0 IGO"
    )
    demo = BriefingPack(
        query="What has WHO done to support rehabilitation and assistive technology in Tajikistan?",
        points=[
            BriefingPoint(
                point=(
                    "WHO supported a pilot initiative in Tajikistan from February 2019 "
                    "to June 2022 that provided over 6,000 assistive products to more "
                    "than 4,000 people at the community level through primary health "
                    "care centres."
                ),
                quote=(
                    "From February 2019 to June 2022, WHO supported a pilot initiative "
                    "called '1stop assistive products service provision in Tajikistan' "
                    "(known as the 1stop project)."
                ),
                citation=_cite,
                iris_url="https://iris.who.int/handle/10665/380686",
                page=85,
                confidence=0.95,
            ),
            BriefingPoint(
                point=(
                    "At the request of Tajikistan's Ministry of Health and Social "
                    "Protection of the Population, WHO provided technical assistance "
                    "for an in-country assessment of rehabilitation."
                ),
                quote=(
                    "At the request of Tajikistan's Ministry of Health and Social "
                    "Protection of the Population (MoHSPP), the WHO Regional Office "
                    "for Europe and the WHO Country Office in Tajikistan provided "
                    "technical assistance for an incountry assessment of rehabilitation "
                    "from 28 November to 16 December 2022."
                ),
                citation=_cite,
                iris_url="https://iris.who.int/handle/10665/380686",
                page=11,
                confidence=0.88,
            ),
        ],
        abstained=False,
    )
    path = render_briefing(demo, "/tmp/who_render_demo.docx")
    print(f"rendered: {path}")
